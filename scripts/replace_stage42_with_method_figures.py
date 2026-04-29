from __future__ import annotations

from pathlib import Path
import shutil

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches


ROOT = Path(__file__).resolve().parents[1]
SPEC_DIR = ROOT / "specs"
FIG_DIR = SPEC_DIR / "figures"

SRC = SPEC_DIR / (
    "\u6bd5\u4e1a\u8bba\u6587_\u5b8c\u6574\u521d\u7a3f_20260423_"
    "4_2\u65b9\u6cd5\u8fc7\u7a0b\u56fe\u7248.docx"
)
OUT = SPEC_DIR / (
    "\u6bd5\u4e1a\u8bba\u6587_\u5b8c\u6574\u521d\u7a3f_20260423_"
    "4_2\u65b9\u6cd5\u8fc7\u7a0b\u56fe\u7248_\u4fee\u6b63.docx"
)

FIGURES = {
    "\u56fe4-2 \u8f6e\u8f8b\u8f6e\u5fc3\u56de\u8f6c\u4f53\u9636\u6bb5\u611f\u77e5\u4e0e\u5efa\u6a21\u793a\u610f\u56fe": FIG_DIR
    / "stage01_method_pipeline_slice.png",
    "\u56fe4-3 PCD \u5b54\u9636\u6bb5\u611f\u77e5\u4e0e\u5efa\u6a21\u793a\u610f\u56fe": FIG_DIR
    / "stage02_method_pipeline_pcd.png",
    "\u56fe4-4 \u8f6e\u5fc3\u975e\u5b54\u7279\u5f81\u9636\u6bb5\u611f\u77e5\u4e0e\u5efa\u6a21\u793a\u610f\u56fe": FIG_DIR
    / "stage03_method_pipeline_nonhole.png",
    "\u56fe4-5 \u8f90\u6761\u751f\u6210\u9636\u6bb5\u611f\u77e5\u4e0e\u5efa\u6a21\u793a\u610f\u56fe": FIG_DIR
    / "stage04_method_pipeline_spoke.png",
}


def clear_paragraph(paragraph) -> None:
    p = paragraph._p
    for child in list(p):
        p.remove(child)


def has_drawing(paragraph) -> bool:
    return bool(paragraph._p.xpath(".//w:drawing"))


def replace_image_before_caption(doc: Document, caption: str, image_path: Path) -> None:
    paragraphs = list(doc.paragraphs)
    for idx, paragraph in enumerate(paragraphs):
        if paragraph.text.strip() != caption:
            continue

        target = None
        for j in range(idx - 1, max(-1, idx - 8), -1):
            if has_drawing(paragraphs[j]) or "\u63d2\u56fe\u4f4d\u7f6e" in paragraphs[j].text:
                target = paragraphs[j]
                break
        if target is None and idx > 0:
            target = paragraphs[idx - 1]
        if target is None:
            raise ValueError(f"no paragraph before caption: {caption}")

        clear_paragraph(target)
        target.alignment = WD_ALIGN_PARAGRAPH.CENTER
        target.paragraph_format.first_line_indent = None
        target.paragraph_format.space_before = None
        target.paragraph_format.space_after = None
        run = target.add_run()
        run.add_picture(str(image_path), width=Inches(6.15))

        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        paragraph.paragraph_format.first_line_indent = None
        paragraph.paragraph_format.space_before = None
        paragraph.paragraph_format.space_after = None
        return

    raise ValueError(f"caption not found: {caption}")


def main() -> None:
    if not SRC.exists():
        raise FileNotFoundError(SRC)
    for image in FIGURES.values():
        if not image.exists():
            raise FileNotFoundError(image)

    shutil.copy2(SRC, OUT)
    doc = Document(OUT)
    for caption, image_path in FIGURES.items():
        replace_image_before_caption(doc, caption, image_path)
    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    main()
